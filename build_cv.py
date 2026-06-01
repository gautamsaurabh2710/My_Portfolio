from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Saurabh_Gautam_Backend_Developer_CV.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
LIGHT_FILL = "F2F4F7"
BORDER = "DADCE0"


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(paragraph, before=0, after=6, line=1.10):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_hyperlink(paragraph, text, url, color="2E74B5"):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def section_heading(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=10, after=4, line=1.10)
    r = p.add_run(text.upper())
    set_run_font(r, size=12.5, color=BLUE, bold=True)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_para_spacing(p, before=0, after=4, line=1.167)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=RGBColor(0, 0, 0))
    return p


def project(doc, title, meta, bullets):
    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=2, line=1.10)
    r = p.add_run(title)
    set_run_font(r, size=11.5, color=NAVY, bold=True)
    r = p.add_run(" | " + meta)
    set_run_font(r, size=10.5, color=MUTED, italic=True)
    for item in bullets:
        bullet(doc, item)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    st = styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.line_spacing = 1.10

title = doc.add_paragraph()
set_para_spacing(title, before=0, after=2, line=1.0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("SAURABH GAUTAM")
set_run_font(r, size=22, color=NAVY, bold=True)

subtitle = doc.add_paragraph()
set_para_spacing(subtitle, before=0, after=6, line=1.0)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Backend Developer | Golang APIs | Authentication | Database-Driven Systems")
set_run_font(r, size=10.8, color=MUTED, bold=True)

contact = doc.add_paragraph()
set_para_spacing(contact, before=0, after=10, line=1.0)
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(contact.add_run("+91 6393235581 | "), size=9.8, color=MUTED)
add_hyperlink(contact, "gautamsaurabh90@gmail.com", "mailto:gautamsaurabh90@gmail.com")
set_run_font(contact.add_run(" | "), size=9.8, color=MUTED)
add_hyperlink(contact, "GitHub", "https://github.com/gautamsaurabh2710")
set_run_font(contact.add_run(" | "), size=9.8, color=MUTED)
add_hyperlink(contact, "LinkedIn", "https://www.linkedin.com/in/saurabh-gautam-9a7697257")

summary_table = doc.add_table(rows=1, cols=1)
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
summary_cell = summary_table.cell(0, 0)
summary_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
shade_cell(summary_cell, LIGHT_FILL)
set_cell_border(summary_cell)
set_cell_margins(summary_cell, top=110, bottom=110, start=160, end=160)
p = summary_cell.paragraphs[0]
set_para_spacing(p, before=0, after=0, line=1.10)
r = p.add_run(
    "Backend developer focused on building secure server-side applications, clean API contracts, "
    "authentication flows, and maintainable data-driven systems. Comfortable with Golang APIs, "
    "REST architecture, login/register workflows, OTP authentication, and responsive HTML/CSS projects."
)
set_run_font(r, size=10.6, color=RGBColor(0, 0, 0))

section_heading(doc, "Technical Skills")
skills = [
    ("Backend", "Golang, Node.js, Express, REST APIs, API design"),
    ("Authentication", "Register/login flows, OTP authentication, JWT concepts, protected routes"),
    ("Data", "SQL, MongoDB, database design, CRUD systems"),
    ("Frontend", "HTML, CSS, responsive landing pages"),
    ("Tools & Concepts", "Docker, GitHub, system design, debugging, problem solving"),
]
skill_table = doc.add_table(rows=0, cols=2)
skill_table.alignment = WD_TABLE_ALIGNMENT.CENTER
skill_table.autofit = False
for label, value in skills:
    row = skill_table.add_row()
    row.cells[0].width = Inches(1.45)
    row.cells[1].width = Inches(5.15)
    for cell in row.cells:
        set_cell_border(cell)
        set_cell_margins(cell, top=80, bottom=80, start=120, end=120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(row.cells[0], LIGHT_FILL)
    p0 = row.cells[0].paragraphs[0]
    set_para_spacing(p0, before=0, after=0, line=1.10)
    set_run_font(p0.add_run(label), size=10.2, color=NAVY, bold=True)
    p1 = row.cells[1].paragraphs[0]
    set_para_spacing(p1, before=0, after=0, line=1.10)
    set_run_font(p1.add_run(value), size=10.2, color=RGBColor(0, 0, 0))

section_heading(doc, "Projects")
project(
    doc,
    "Password Manager API",
    "Golang, REST APIs, OTP Authentication",
    [
        "Built a secure password manager backend with register and login flows.",
        "Implemented OTP authentication and protected user vault functionality.",
        "Designed API endpoints for authentication, user access, and secure password-management features.",
    ],
)
project(
    doc,
    "Backend Developer Portfolio",
    "HTML, CSS, responsive portfolio",
    [
        "Created a responsive personal portfolio with a backend developer theme.",
        "Showcased backend skills, real project work, contact links, and relative project images.",
    ],
)
project(
    doc,
    "Hospital Landing Pages",
    "HTML, CSS, healthcare landing pages",
    [
        "Built hospital landing pages inspired by Fortis Hospital and Max Hospital.",
        "Designed clean healthcare-focused sections for responsive presentation and user-friendly layout.",
    ],
)

section_heading(doc, "Education")
bullet(doc, "M.C.A., Babasaheb Bhimrao Ambedkar University, 2024, and Lucknow.")

section_heading(doc, "Profile Links")
p = doc.add_paragraph()
set_para_spacing(p, before=0, after=3, line=1.10)
set_run_font(p.add_run("GitHub: "), size=10.5, color=NAVY, bold=True)
add_hyperlink(p, "github.com/gautamsaurabh2710", "https://github.com/gautamsaurabh2710")
p = doc.add_paragraph()
set_para_spacing(p, before=0, after=3, line=1.10)
set_run_font(p.add_run("LinkedIn: "), size=10.5, color=NAVY, bold=True)
add_hyperlink(p, "linkedin.com/in/saurabh-gautam-9a7697257", "https://www.linkedin.com/in/saurabh-gautam-9a7697257")
p = doc.add_paragraph()
set_para_spacing(p, before=0, after=3, line=1.10)
set_run_font(p.add_run("Naukri: "), size=10.5, color=NAVY, bold=True)
set_run_font(p.add_run("Add your Naukri profile URL here."), size=10.5, color=MUTED, italic=True)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(footer, before=0, after=0, line=1.0)
set_run_font(footer.add_run("Saurabh Gautam - Backend Developer CV"), size=9, color=MUTED)

doc.save(OUT)
print(OUT)
